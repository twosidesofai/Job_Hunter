"""
Exporter agent that generates PDF/DOCX files using OpenAI agent structure.
"""
import os
import json
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
from datetime import datetime

# Optional imports with fallbacks
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

from base_agent import BaseAgent, AgentResponse
from models import TailoredResume, CoverLetter, JobPosting, ExportFormat
from config import Config

class ExporterAgent(BaseAgent):
    """
    Agent responsible for exporting resumes and cover letters to PDF/DOCX formats.
    Uses OpenAI to help with formatting decisions and content optimization.
    """
    
    def __init__(self):
        system_prompt = """
        You are a document formatting and export specialist. Your role is to:
        
        1. Format resumes and cover letters for professional presentation
        2. Optimize content layout for ATS compatibility and visual appeal
        3. Ensure proper document structure and formatting
        4. Generate clean, professional PDF and DOCX files
        5. Manage file naming conventions and organization
        
        When formatting documents:
        - Use professional fonts and spacing
        - Ensure proper hierarchy with headings and sections
        - Maintain consistent formatting throughout
        - Optimize for both digital viewing and printing
        - Ensure ATS-friendly formatting (no images, proper text structure)
        - Use appropriate file naming conventions
        
        Focus on creating documents that are both visually appealing and functionally effective.
        """
        super().__init__("Exporter", system_prompt)
        self.supported_functions = ["format_resume", "format_cover_letter", "optimize_layout"]
        
        # Ensure Applications directory exists
        os.makedirs(Config.APPLICATIONS_DIR, exist_ok=True)
    
    def execute_function(self, function_name: str, args: Dict[str, Any]) -> Any:
        """Execute specific export functions."""
        if function_name == "format_resume":
            return self._format_resume_content(args.get("resume", {}))
        elif function_name == "format_cover_letter":
            return self._format_cover_letter_content(args.get("cover_letter", {}))
        elif function_name == "optimize_layout":
            return self._optimize_document_layout(args.get("content", ""), args.get("document_type", ""))
        else:
            raise ValueError(f"Unknown function: {function_name}")
    
    def process_request(self, request: str, context: Optional[Dict[str, Any]] = None) -> AgentResponse:
        """Process export requests."""
        try:
            if "export" in request.lower():
                if context:
                    job_posting = JobPosting(**context["job_posting"]) if "job_posting" in context else None
                    resume = TailoredResume(**context["resume"]) if "resume" in context else None
                    cover_letter = CoverLetter(**context["cover_letter"]) if "cover_letter" in context else None
                    export_format = context.get("format", ExportFormat.BOTH)
                    
                    if job_posting and (resume or cover_letter):
                        file_paths = self.export_application_documents(
                            job_posting=job_posting,
                            resume=resume,
                            cover_letter=cover_letter,
                            export_format=export_format
                        )
                        
                        return AgentResponse(
                            success=True,
                            data={"file_paths": file_paths},
                            message=f"Successfully exported documents to {len(file_paths)} files"
                        )
                    else:
                        return AgentResponse(
                            success=False,
                            message="Job posting and at least one document (resume or cover letter) required",
                            errors=["Missing required context data"]
                        )
            else:
                # Use OpenAI for general formatting questions
                messages = self._prepare_messages(request, context)
                response = self._make_openai_request(messages)
                return self._handle_response(response)
                
        except Exception as e:
            self.logger.error(f"Error processing export request: {str(e)}")
            return AgentResponse(
                success=False,
                message="Failed to process export request",
                errors=[str(e)]
            )
    
    def export_application_documents(
        self,
        job_posting: JobPosting,
        resume: Optional[TailoredResume] = None,
        cover_letter: Optional[CoverLetter] = None,
        export_format: ExportFormat = ExportFormat.BOTH
    ) -> Dict[str, str]:
        """
        Export resume and/or cover letter to specified formats.
        
        Args:
            job_posting: Job posting information for file naming
            resume: Tailored resume to export
            cover_letter: Cover letter to export
            export_format: Export format (PDF, DOCX, or BOTH)
            
        Returns:
            Dictionary mapping document types to file paths
        """
        file_paths = {}
        
        try:
            # Create company-specific folder
            company_folder = self._create_company_folder(job_posting)
            
            # Export resume if provided
            if resume:
                resume_paths = self._export_resume(resume, job_posting, company_folder, export_format)
                file_paths.update(resume_paths)
            
            # Export cover letter if provided
            if cover_letter:
                cover_letter_paths = self._export_cover_letter(cover_letter, job_posting, company_folder, export_format)
                file_paths.update(cover_letter_paths)
            
            self.logger.info(f"Successfully exported {len(file_paths)} files for {job_posting.company} - {job_posting.title}")
            return file_paths
            
        except Exception as e:
            self.logger.error(f"Error exporting documents: {str(e)}")
            raise
    
    def _create_company_folder(self, job_posting: JobPosting) -> str:
        """
        Create a folder for the company and job application.
        
        Args:
            job_posting: Job posting information
            
        Returns:
            Path to the created folder
        """
        # Sanitize company name and job title for file system
        company_clean = self._sanitize_filename(job_posting.company)
        job_title_clean = self._sanitize_filename(job_posting.title)
        
        folder_name = f"{company_clean}_{job_title_clean}"
        folder_path = os.path.join(Config.APPLICATIONS_DIR, folder_name)
        
        os.makedirs(folder_path, exist_ok=True)
        return folder_path
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for file system compatibility."""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length and remove extra spaces
        filename = filename.strip()[:50]
        filename = '_'.join(filename.split())
        
        return filename
    
    def _export_resume(
        self, 
        resume: TailoredResume, 
        job_posting: JobPosting, 
        folder_path: str, 
        export_format: ExportFormat
    ) -> Dict[str, str]:
        """Export resume to specified format(s)."""
        file_paths = {}
        base_filename = f"Resume_{self._sanitize_filename(job_posting.company)}_{self._sanitize_filename(job_posting.title)}"
        
        try:
            if export_format in [ExportFormat.PDF, ExportFormat.BOTH]:
                pdf_path = os.path.join(folder_path, f"{base_filename}.pdf")
                self._create_resume_pdf(resume, pdf_path)
                file_paths["resume_pdf"] = pdf_path
            
            if export_format in [ExportFormat.DOCX, ExportFormat.BOTH]:
                docx_path = os.path.join(folder_path, f"{base_filename}.docx")
                self._create_resume_docx(resume, docx_path)
                file_paths["resume_docx"] = docx_path
            
            return file_paths
            
        except Exception as e:
            self.logger.error(f"Error exporting resume: {str(e)}")
            raise
    
    def _export_cover_letter(
        self, 
        cover_letter: CoverLetter, 
        job_posting: JobPosting, 
        folder_path: str, 
        export_format: ExportFormat
    ) -> Dict[str, str]:
        """Export cover letter to specified format(s)."""
        file_paths = {}
        base_filename = f"CoverLetter_{self._sanitize_filename(job_posting.company)}_{self._sanitize_filename(job_posting.title)}"
        
        try:
            if export_format in [ExportFormat.PDF, ExportFormat.BOTH]:
                pdf_path = os.path.join(folder_path, f"{base_filename}.pdf")
                self._create_cover_letter_pdf(cover_letter, pdf_path)
                file_paths["cover_letter_pdf"] = pdf_path
            
            if export_format in [ExportFormat.DOCX, ExportFormat.BOTH]:
                docx_path = os.path.join(folder_path, f"{base_filename}.docx")
                self._create_cover_letter_docx(cover_letter, docx_path)
                file_paths["cover_letter_docx"] = docx_path
            
            return file_paths
            
        except Exception as e:
            self.logger.error(f"Error exporting cover letter: {str(e)}")
            raise
    
    def _create_resume_pdf(self, resume: TailoredResume, file_path: str):
        """Create PDF version of resume."""
        if not FPDF_AVAILABLE:
            self.logger.error("FPDF not available, cannot create PDF")
            # Create a placeholder text file instead
            with open(file_path.replace('.pdf', '.txt'), 'w') as f:
                f.write(f"Resume for {resume.personal_info.get('name', 'Unknown')}\n")
                f.write(f"Summary: {resume.summary}\n")
            return
            
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Personal Info
        personal_info = resume.personal_info
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, personal_info.get("name", ""), ln=True, align="C")
        
        pdf.set_font("Arial", size=10)
        contact_line = f"{personal_info.get('email', '')} | {personal_info.get('phone', '')} | {personal_info.get('location', '')}"
        pdf.cell(0, 8, contact_line, ln=True, align="C")
        pdf.ln(5)
        
        # Professional Summary
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "PROFESSIONAL SUMMARY", ln=True)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        
        pdf.set_font("Arial", size=10)
        summary_lines = self._wrap_text(resume.summary, 90)
        for line in summary_lines:
            pdf.cell(0, 6, line, ln=True)
        pdf.ln(3)
        
        # Experience
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "PROFESSIONAL EXPERIENCE", ln=True)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        
        for exp in resume.experience:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, f"{exp.title} - {exp.company}", ln=True)
            
            pdf.set_font("Arial", "I", 9)
            pdf.cell(0, 5, exp.duration, ln=True)
            
            pdf.set_font("Arial", size=9)
            desc_lines = self._wrap_text(exp.description, 100)
            for line in desc_lines:
                pdf.cell(0, 5, line, ln=True)
            pdf.ln(2)
        
        # Skills
        if resume.skills:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "TECHNICAL SKILLS", ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            pdf.set_font("Arial", size=10)
            skills_text = " • ".join(resume.skills)
            skills_lines = self._wrap_text(skills_text, 90)
            for line in skills_lines:
                pdf.cell(0, 6, line, ln=True)
            pdf.ln(3)
        
        # Education
        if resume.education:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "EDUCATION", ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(2)
            
            for edu in resume.education:
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 6, f"{edu.degree} - {edu.institution}", ln=True)
                if edu.year:
                    pdf.set_font("Arial", "I", 9)
                    pdf.cell(0, 5, edu.year, ln=True)
                pdf.ln(1)
        
        pdf.output(file_path)
    
    def _create_resume_docx(self, resume: TailoredResume, file_path: str):
        """Create DOCX version of resume."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available, cannot create DOCX")
            # Create a placeholder text file instead
            with open(file_path.replace('.docx', '.txt'), 'w') as f:
                f.write(f"Resume for {resume.personal_info.get('name', 'Unknown')}\n")
                f.write(f"Summary: {resume.summary}\n")
            return
            
        doc = Document()
        
        # Personal Info
        personal_info = resume.personal_info
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.bold = True
        name_run.font.size = Inches(0.2)
        
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_line = f"{personal_info.get('email', '')} | {personal_info.get('phone', '')} | {personal_info.get('location', '')}"
        contact_para.add_run(contact_line)
        
        # Professional Summary
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        doc.add_paragraph(resume.summary)
        
        # Experience
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
        for exp in resume.experience:
            exp_para = doc.add_paragraph()
            title_run = exp_para.add_run(f"{exp.title} - {exp.company}")
            title_run.bold = True
            
            duration_para = doc.add_paragraph(exp.duration)
            duration_para.paragraph_format.left_indent = Inches(0.5)
            
            desc_para = doc.add_paragraph(exp.description)
            desc_para.paragraph_format.left_indent = Inches(0.5)
        
        # Skills
        if resume.skills:
            doc.add_heading("TECHNICAL SKILLS", level=2)
            skills_text = " • ".join(resume.skills)
            doc.add_paragraph(skills_text)
        
        # Education
        if resume.education:
            doc.add_heading("EDUCATION", level=2)
            for edu in resume.education:
                edu_para = doc.add_paragraph()
                degree_run = edu_para.add_run(f"{edu.degree} - {edu.institution}")
                degree_run.bold = True
                if edu.year:
                    edu_para.add_run(f" ({edu.year})")
        
        doc.save(file_path)
    
    def _create_cover_letter_pdf(self, cover_letter: CoverLetter, file_path: str):
        """Create PDF version of cover letter."""
        if not FPDF_AVAILABLE:
            self.logger.error("FPDF not available, cannot create PDF")
            # Create a placeholder text file instead
            with open(file_path.replace('.pdf', '.txt'), 'w') as f:
                f.write(cover_letter.content)
            return
            
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        
        # Date
        pdf.cell(0, 10, datetime.now().strftime("%B %d, %Y"), ln=True, align="R")
        pdf.ln(10)
        
        # Cover letter content
        content_lines = cover_letter.content.split('\n')
        for line in content_lines:
            if line.strip():
                wrapped_lines = self._wrap_text(line, 85)
                for wrapped_line in wrapped_lines:
                    pdf.cell(0, 6, wrapped_line, ln=True)
            else:
                pdf.ln(4)
        
        pdf.output(file_path)
    
    def _create_cover_letter_docx(self, cover_letter: CoverLetter, file_path: str):
        """Create DOCX version of cover letter."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available, cannot create DOCX")
            # Create a placeholder text file instead
            with open(file_path.replace('.docx', '.txt'), 'w') as f:
                f.write(cover_letter.content)
            return
            
        doc = Document()
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        
        doc.add_paragraph()  # Empty line
        
        # Cover letter content
        content_paragraphs = cover_letter.content.split('\n\n')
        for para_text in content_paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        doc.save(file_path)
    
    def _wrap_text(self, text: str, max_length: int) -> list[str]:
        """Wrap text to specified length."""
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            if len(current_line + " " + word) <= max_length:
                current_line += " " + word if current_line else word
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _format_resume_content(self, resume: Dict[str, Any]) -> Dict[str, Any]:
        """Format resume content for better presentation."""
        return {"formatted": True, "content": resume}
    
    def _format_cover_letter_content(self, cover_letter: Dict[str, Any]) -> Dict[str, Any]:
        """Format cover letter content for better presentation."""
        return {"formatted": True, "content": cover_letter}
    
    def _optimize_document_layout(self, content: str, document_type: str) -> Dict[str, Any]:
        """Optimize document layout."""
        return {"optimized": True, "suggestions": []}